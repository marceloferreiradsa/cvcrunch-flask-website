import docx
import os
import re
import uuid
import time
import io

from pdfminer.high_level import extract_pages
from pdfminer.layout import LTTextContainer, LAParams

IMG_ALLOWED_EXTENSIONS = set(['jpg', 'png', 'webp'])
ALLOWED_EXTENSIONS = set(['txt', 'pdf', 'docx'])
  
def allowed_file(filename):
 return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def parse_file(file, filename):
    print(file, filename)
    new_text = ""
    
    if filename.rsplit('.', 1)[1] == 'txt':
        pass 
    elif filename.rsplit('.', 1)[1] == 'docx':
        doc = docx.Document(file)
        new_text = ''
        for p in doc.paragraphs:
            new_text+=p.text.replace('\t', ' ') + ' '
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        new_text += paragraph.text + ' '
        
    elif filename.rsplit('.', 1)[1] == 'pdf':
        extracted_text = []
        file_path = os.path.join(os.path.dirname(__file__), '..', 'static', 'uploads', filename)

        for page_layout in extract_pages(os.path.join(file_path)):
            for element in page_layout:
                if isinstance(element, LTTextContainer):
                    lines = element.get_text().split('\n')
                    for line in lines:
                        line = line.strip()
                        if line:  # Skip empty lines
                            extracted_text.append(line)
        new_text = ' '.join(extracted_text)
            
    new_text = re.sub(' +', ' ', new_text.strip())
    return new_text
        
